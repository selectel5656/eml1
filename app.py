import os
import base64
import random
import string
import time
import requests
import quopri
import threading
import json
from PIL import Image, ImageDraw
from functools import wraps
from flask import (
    Flask, render_template, request, redirect, url_for, session, flash,
    send_from_directory
)
from werkzeug.utils import secure_filename
from models import db, User, EmailEntry, Macro, Attachment, Proxy, ApiAccount, Setting
from api_client import ApiClient

app = Flask(__name__)
app.config['SECRET_KEY'] = 'dev'
app.config['UPLOAD_FOLDER'] = os.path.join(app.root_path, 'uploads')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

CONFIG_FILE = os.path.join(app.root_path, 'config.json')
db_configured = False


def configure_db(uri: str) -> None:
    app.config['SQLALCHEMY_DATABASE_URI'] = uri
    db.init_app(app)
    init_db()


def init_db():
    """Initialize database and seed default settings."""
    with app.app_context():
        db.create_all()
        if not User.query.filter_by(username='admin').first():
            u = User(username='admin')
            u.set_password('admin')
            db.session.add(u)
        if not Setting.query.filter_by(key='domain').first():
            db.session.add(Setting(key='domain', value='domen.ru'))
        if not Setting.query.filter_by(key='user_agent').first():
            db.session.add(Setting(key='user_agent', value=ApiClient.USER_AGENT))
        if not Setting.query.filter_by(key='per_account_limit').first():
            db.session.add(Setting(key='per_account_limit', value='1'))
        if not Setting.query.filter_by(key='cycle_accounts').first():
            db.session.add(Setting(key='cycle_accounts', value='no'))
        if not Setting.query.filter_by(key='send_attempts').first():
            db.session.add(Setting(key='send_attempts', value='1'))
        if not Setting.query.filter_by(key='server_timeout').first():
            db.session.add(Setting(key='server_timeout', value='30'))
        if not Setting.query.filter_by(key='pause_between').first():
            db.session.add(Setting(key='pause_between', value='0'))
        if not Setting.query.filter_by(key='recipients_per_message').first():
            db.session.add(Setting(key='recipients_per_message', value='1'))
        if not Setting.query.filter_by(key='recipient_method').first():
            db.session.add(Setting(key='recipient_method', value='bcc'))
        if not Setting.query.filter_by(key='first_recipient_to').first():
            db.session.add(Setting(key='first_recipient_to', value='no'))
        if not Setting.query.filter_by(key='quality_every').first():
            db.session.add(Setting(key='quality_every', value='0'))
        if not Setting.query.filter_by(key='quality_email').first():
            db.session.add(Setting(key='quality_email', value=''))
        if not Setting.query.filter_by(key='threads').first():
            db.session.add(Setting(key='threads', value='1'))
        if not Setting.query.filter_by(key='total_sent').first():
            db.session.add(Setting(key='total_sent', value='0'))
        db.session.commit()
        ApiAccount.query.update({ApiAccount.in_use: False})
        db.session.commit()


if os.path.exists(CONFIG_FILE):
    with open(CONFIG_FILE) as f:
        data = json.load(f)
    configure_db(data.get('database_uri', 'sqlite:///data.db'))
    db_configured = True


@app.before_request
def ensure_setup():
    if not db_configured and request.endpoint not in ('setup', 'static'):
        return redirect(url_for('setup'))


@app.route('/setup', methods=['GET', 'POST'])
def setup():
    global db_configured
    if db_configured:
        return redirect(url_for('login'))
    if request.method == 'POST':
        db_type = request.form['db_type']
        if db_type == 'sqlite':
            uri = 'sqlite:///data.db'
        else:
            host = request.form.get('host', 'localhost')
            port = request.form.get('port', '3306')
            user = request.form.get('user', '')
            password = request.form.get('password', '')
            name = request.form.get('name', '')
            uri = f"mysql+pymysql://{user}:{password}@{host}:{port}/{name}"
        with open(CONFIG_FILE, 'w') as f:
            json.dump({'database_uri': uri}, f)
        configure_db(uri)
        db_configured = True
        return redirect(url_for('login'))
    return render_template('setup.html')


send_threads: list[threading.Thread] = []
sending = False
stop_flag = False

def login_required(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return wrapper


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        user = User.query.filter_by(username=request.form['username']).first()
        if user and user.check_password(request.form['password']):
            session['user_id'] = user.id
            return redirect(url_for('letter'))
        flash('Неверный логин или пароль')
    return render_template('login.html')


@app.route('/logout')
def logout():
    session.pop('user_id', None)
    return redirect(url_for('login'))


@app.route('/')
@login_required
def index():
    return redirect(url_for('letter'))


# ------- Letter -------


def get_setting(key: str, default: str = '') -> str:
    setting = Setting.query.filter_by(key=key).first()
    return setting.value if setting else default


def get_domain() -> str:
    return get_setting('domain', 'domen.ru')


def get_user_agent() -> str:
    return get_setting('user_agent', ApiClient.USER_AGENT)


def check_proxy(address: str) -> bool:
    try:
        proxies = {'http': f'http://{address}', 'https': f'http://{address}'}
        r = requests.get('https://httpbin.org/ip', proxies=proxies, timeout=5)
        return r.status_code == 200
    except Exception:
        return False


def acquire_proxy() -> str | None:
    for proxy in Proxy.query.filter_by(in_use=False).all():
        if check_proxy(proxy.address):
            proxy.in_use = True
            db.session.commit()
            return proxy.address
    Proxy.query.update({Proxy.in_use: False})
    db.session.commit()
    return None


def release_proxy(account: ApiAccount) -> None:
    if account.proxy_id:
        proxy = Proxy.query.get(account.proxy_id)
        if proxy:
            proxy.in_use = False
        account.proxy_id = None
        db.session.commit()


def acquire_proxy_for_account(account: ApiAccount) -> str | None:
    if account.proxy_id:
        proxy = Proxy.query.get(account.proxy_id)
        if proxy and check_proxy(proxy.address):
            return proxy.address
        release_proxy(account)
    addr = acquire_proxy()
    if addr:
        proxy = Proxy.query.filter_by(address=addr).first()
        if proxy:
            account.proxy_id = proxy.id
            db.session.commit()
    return addr


def send_batch(subject_raw: str, body_raw: str, selected: list[str]) -> bool:
    """Send a single batch of emails according to current settings."""
    limit = int(get_setting('per_account_limit', '1'))
    cycle = get_setting('cycle_accounts', 'no') == 'yes'
    attempts = int(get_setting('send_attempts', '1'))
    timeout = int(get_setting('server_timeout', '30'))
    pause = int(get_setting('pause_between', '0'))
    rec_count = int(get_setting('recipients_per_message', '1'))
    method = get_setting('recipient_method', 'bcc')
    first = get_setting('first_recipient_to', 'no') == 'yes'
    q_every = int(get_setting('quality_every', '0'))
    q_email = get_setting('quality_email', '')

    account = ApiAccount.query.filter(
        ApiAccount.send_count < limit, ApiAccount.in_use.is_(False)
    ).order_by(ApiAccount.id).first()
    if not account and cycle:
        ApiAccount.query.update({ApiAccount.send_count: 0, ApiAccount.in_use: False})
        db.session.commit()
        account = ApiAccount.query.filter(
            ApiAccount.send_count < limit, ApiAccount.in_use.is_(False)
        ).order_by(ApiAccount.id).first()
    if not account:
        return False
    account.in_use = True
    db.session.commit()

    success = False
    try:
        for _proxy_try in range(3):
            proxy_addr = acquire_proxy_for_account(account)
            client = ApiClient(
                get_domain(),
                account.api_key,
                account.uuid,
                login=account.login,
                from_name=account.first_name or '',
                user_agent=get_user_agent(),
                proxy=proxy_addr,
                timeout=timeout,
            )
            if not client.check_account():
                release_proxy(account)
                continue

            att_ids: list[str] = []
            for sid in selected:
                att = Attachment.query.get(int(sid))
                if att:
                    if (not att.inline) or att.upload_to_server:
                        if not att.remote_id:
                            send_name = att.send_filename or att.filename
                            res = client.upload_attachment(att.path, send_name)
                            att.remote_id = res.get('id')
                            att.remote_url = res.get('url')
                            db.session.commit()
                        if att.remote_id:
                            att_ids.append(att.remote_id)

            subject = render_macros(subject_raw)
            body = render_macros(body_raw)
            emails = EmailEntry.query.filter_by(sent=False, in_progress=False).limit(rec_count).all()
            if not emails:
                return False
            for e in emails:
                e.in_progress = True
            db.session.commit()
            recipients = [e.email for e in emails]
            for _ in range(attempts):
                operation_id = client.generate_operation_id()
                if not operation_id:
                    continue
                if client.send_mail(
                    subject,
                    body,
                    recipients,
                    att_ids,
                    operation_id,
                    method=method,
                    first_to=first,
                ):
                    success = True
                    break
            if success:
                break
            release_proxy(account)
    finally:
        account.in_use = False
        db.session.commit()

    if success:
        account.send_count = account.send_count + 1
        total_sent = int(get_setting('total_sent', '0')) + 1
        Setting.query.filter_by(key='total_sent').first().value = str(total_sent)
        for e in emails:
            e.sent = True
            e.in_progress = False
        db.session.commit()
        if q_every and q_email and total_sent % q_every == 0:
            qc_id = client.generate_operation_id()
            if qc_id:
                client.send_mail(
                    subject,
                    body,
                    [q_email],
                    att_ids,
                    qc_id,
                    method='to',
                )
        if pause > 0:
            time.sleep(pause)
        return True
    else:
        for e in emails:
            e.in_progress = False
        db.session.commit()
        return False


def send_worker(subject_raw: str, body_raw: str, selected: list[str]) -> None:
    """Thread worker that keeps sending batches until stopped."""
    with app.app_context():
        while not stop_flag:
            if not EmailEntry.query.filter_by(sent=False, in_progress=False).first():
                break
            if not send_batch(subject_raw, body_raw, selected):
                time.sleep(1)

def evaluate_macro(m: Macro) -> str:
    """Evaluate macro value and update usage counters."""
    if m.frequency > 1 and m.usage_count and m.usage_count % m.frequency != 0 and m.current_value:
        m.usage_count += 1
        db.session.commit()
        return m.current_value

    cfg = m.config or {}
    value = ''
    if m.macro_type == 'counter':
        start = int(cfg.get('start', 0))
        step = int(cfg.get('step', 1))
        current = int(cfg.get('current', start))
        value = str(current)
        cfg['current'] = current + step
    elif m.macro_type == 'random':
        chars = cfg.get('chars', string.ascii_letters)
        min_len = int(cfg.get('min_len', 5))
        max_len = int(cfg.get('max_len', 10))
        n = random.randint(min_len, max_len)
        value = ''.join(random.choice(chars) for _ in range(n))
    elif m.macro_type == 'list':
        items = cfg.get('items', [])
        if items:
            mode = cfg.get('mode', 'random')
            if cfg.get('words_min'):
                wmin = int(cfg.get('words_min', 1))
                wmax = int(cfg.get('words_max', wmin))
                smin = int(cfg.get('sent_min', 1))
                smax = int(cfg.get('sent_max', smin))
                pmin = int(cfg.get('para_min', 1))
                pmax = int(cfg.get('para_max', pmin))
                as_html = cfg.get('as_html', False)
                extra = cfg.get('html_extra', '')
                paragraphs = []
                for _ in range(random.randint(pmin, pmax)):
                    sentences = []
                    for _ in range(random.randint(smin, smax)):
                        count = random.randint(wmin, wmax)
                        words = []
                        if mode == 'sequential':
                            idx = int(cfg.get('index', 0))
                            for _ in range(count):
                                words.append(items[idx % len(items)])
                                idx += 1
                            cfg['index'] = idx
                        else:
                            words = [random.choice(items) for _ in range(count)]
                        sentence = ' '.join(words).strip()
                        if sentence:
                            sentence = sentence[0].upper() + sentence[1:]
                        sentences.append(sentence + '.')
                    paragraph = ' '.join(sentences)
                    if as_html:
                        paragraphs.append(f'<p>{paragraph}</p>')
                    else:
                        paragraphs.append(paragraph)
                value = '\n\n'.join(paragraphs)
                if as_html and extra:
                    value += extra
            else:
                if mode == 'sequential':
                    idx = int(cfg.get('index', 0))
                    value = items[idx % len(items)]
                    cfg['index'] = idx + 1
                else:
                    value = random.choice(items)
        else:
            value = ''
    elif m.macro_type == 'multi':
        expr = cfg.get('expr', '')
        encoding = cfg.get('encoding', 'none')
        value = expr
        for other in Macro.query.all():
            if other.id == m.id:
                continue
            val = evaluate_macro(other)
            value = value.replace(f'{{$' + other.name + '}}', val)
        if encoding == 'base64':
            value = base64.b64encode(value.encode()).decode()
        elif encoding == 'quoted-printable':
            value = quopri.encodestring(value.encode()).decode()
    else:
        value = cfg.get('value', '')
    m.current_value = value
    m.usage_count = 1
    m.config = cfg
    db.session.commit()
    return value


def preview_macro_value(macro_type: str, cfg: dict) -> str:
    """Generate a sample value for macro creation preview."""
    if macro_type == 'counter':
        start = int(cfg.get('start', 0))
        step = int(cfg.get('step', 1))
        current = int(cfg.get('current', start))
        return str(current)
    if macro_type == 'random':
        chars = cfg.get('chars', string.ascii_letters)
        min_len = int(cfg.get('min_len', 5))
        max_len = int(cfg.get('max_len', 10))
        n = random.randint(min_len, max_len)
        return ''.join(random.choice(chars) for _ in range(n))
    if macro_type == 'list':
        items = cfg.get('items', [])
        if not items:
            return ''
        mode = cfg.get('mode', 'random')
        if cfg.get('words_min'):
            wmin = int(cfg.get('words_min', 1))
            wmax = int(cfg.get('words_max', wmin))
            smin = int(cfg.get('sent_min', 1))
            smax = int(cfg.get('sent_max', smin))
            pmin = int(cfg.get('para_min', 1))
            pmax = int(cfg.get('para_max', pmin))
            as_html = cfg.get('as_html', False)
            extra = cfg.get('html_extra', '')
            paragraphs = []
            tmp_idx = int(cfg.get('index', 0))
            for _ in range(random.randint(pmin, pmax)):
                sentences = []
                for _ in range(random.randint(smin, smax)):
                    count = random.randint(wmin, wmax)
                    if mode == 'sequential':
                        words = [items[(tmp_idx + i) % len(items)] for i in range(count)]
                        tmp_idx += count
                    else:
                        words = [random.choice(items) for _ in range(count)]
                    sentence = ' '.join(words).strip()
                    if sentence:
                        sentence = sentence[0].upper() + sentence[1:]
                    sentences.append(sentence + '.')
                paragraph = ' '.join(sentences)
                if as_html:
                    paragraphs.append(f'<p>{paragraph}</p>')
                else:
                    paragraphs.append(paragraph)
            value = '\n\n'.join(paragraphs)
            if as_html and extra:
                value += extra
            return value
        if mode == 'sequential':
            idx = int(cfg.get('index', 0))
            return items[idx % len(items)]
        return random.choice(items)
    if macro_type == 'multi':
        expr = cfg.get('expr', '')
        encoding = cfg.get('encoding', 'none')
        value = expr
        for other in Macro.query.all():
            val = evaluate_macro(other)
            value = value.replace(f'{{$' + other.name + '}}', val)
        if encoding == 'base64':
            value = base64.b64encode(value.encode()).decode()
        elif encoding == 'quoted-printable':
            value = quopri.encodestring(value.encode()).decode()
        return value
    return cfg.get('value', '')


def render_macros(text: str) -> str:
    """Replace macro placeholders with values."""
    macros = Macro.query.all()
    for m in macros:
        value = evaluate_macro(m)
        text = text.replace(f'{{$' + m.name + '}}', value)
    attachments = Attachment.query.all()
    for a in attachments:
        if a.macro_url:
            url = a.remote_url or url_for('uploaded_file', filename=a.filename, _external=True)
            text = text.replace(f'{{$' + a.macro_url + '}}', url)
        if a.macro_id:
            text = text.replace(f'{{$' + a.macro_id + '}}', a.remote_id or '')
        if a.macro_base64:
            with open(a.path, 'rb') as f:
                b64 = base64.b64encode(f.read()).decode()
            text = text.replace(f'{{$' + a.macro_base64 + '}}', b64)
    return text


@app.route('/letter', methods=['GET', 'POST'])
@login_required
def letter():
    global sending, stop_flag, send_threads
    attachments = Attachment.query.all()
    macros = Macro.query.all()
    if request.method == 'POST':
        action = request.form.get('action')
        if action == 'start' and not sending:
            subject_raw = request.form.get('subject') or ''
            body_raw = request.form.get('body') or ''
            selected = request.form.getlist('attachments')
            stop_flag = False
            sending = True
            send_threads = []
            for _ in range(int(get_setting('threads', '1'))):
                t = threading.Thread(target=send_worker, args=(subject_raw, body_raw, selected))
                t.start()
                send_threads.append(t)
            flash('Отправка запущена')
        elif action == 'stop' and sending:
            stop_flag = True
            for t in send_threads:
                t.join()
            send_threads = []
            sending = False
            flash('Отправка остановлена')
    return render_template('letter.html', attachments=attachments, macros=macros, sending=sending)


@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


# ------- Email Base -------
import re

@app.route('/email_base', methods=['GET', 'POST'])
@login_required
def email_base():
    if request.method == 'POST':
        fmt = request.form.get('format')
        file = request.files.get('file')
        if file:
            lines = file.read().decode('utf-8').splitlines()
            for line in lines:
                line = line.strip().strip(';')
                if not line:
                    continue
                name = ''
                email = ''
                m = re.match(r'(.*)<([^>]+)>', line)
                if m:
                    name = m.group(1).strip().strip('"')
                    email = m.group(2).strip()
                else:
                    email = line
                if not name:
                    name = email.split('@')[0]
                if not EmailEntry.query.filter_by(email=email).first():
                    db.session.add(EmailEntry(name=name, email=email))
            db.session.commit()
            flash('База загружена')
    emails = EmailEntry.query.limit(50).all()
    count = EmailEntry.query.count()
    return render_template('email_base.html', emails=emails, count=count)


@app.route('/email_base/delete/<int:email_id>')
@login_required
def delete_email(email_id):
    entry = EmailEntry.query.get_or_404(email_id)
    db.session.delete(entry)
    db.session.commit()
    flash('Адрес удален')
    return redirect(url_for('email_base'))


@app.route('/email_base/reset_flags')
@login_required
def reset_email_flags():
    EmailEntry.query.update({EmailEntry.sent: False, EmailEntry.in_progress: False})
    db.session.commit()
    flash('Метки сброшены')
    return redirect(url_for('email_base'))


# ------- Macros -------
@app.route('/macros', methods=['GET', 'POST'])
@login_required
def macros():
    if request.method == 'POST':
        name = request.form.get('name')
        macro_type = request.form.get('macro_type')
        frequency = int(request.form.get('frequency') or 1)
        cfg = {}
        if macro_type == 'counter':
            cfg = {
                'start': int(request.form.get('start') or 0),
                'step': int(request.form.get('step') or 1),
                'current': int(request.form.get('start') or 0)
            }
        elif macro_type == 'random':
            cfg = {
                'chars': request.form.get('chars') or string.ascii_letters,
                'min_len': int(request.form.get('min_len') or 5),
                'max_len': int(request.form.get('max_len') or 10)
            }
        elif macro_type == 'list':
            uploaded = request.files.get('file')
            items = []
            if uploaded:
                text = uploaded.read().decode('utf-8')
                src = request.form.get('source') or 'lines'
                if src == 'words':
                    items = re.findall(r'\w+', text)
                elif src == 'sentences':
                    items = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
                else:
                    items = [line.strip() for line in text.splitlines() if line.strip()]
            cfg = {
                'items': items,
                'mode': request.form.get('mode') or 'random',
                'index': 0,
                'source': request.form.get('source') or 'lines',
            }
            wmin = request.form.get('words_min')
            if wmin:
                cfg.update({
                    'words_min': int(wmin or 1),
                    'words_max': int(request.form.get('words_max') or wmin or 1),
                    'sent_min': int(request.form.get('sent_min') or 1),
                    'sent_max': int(request.form.get('sent_max') or 1),
                    'para_min': int(request.form.get('para_min') or 1),
                    'para_max': int(request.form.get('para_max') or 1),
                    'as_html': bool(request.form.get('as_html')),
                    'html_extra': request.form.get('html_extra') or '',
                })
        elif macro_type == 'multi':
            cfg = {
                'expr': request.form.get('expression') or '',
                'encoding': request.form.get('encoding') or 'none',
            }
        if name:
            db.session.add(Macro(name=name, macro_type=macro_type, config=cfg, frequency=frequency))
            db.session.commit()
            flash('Макрос добавлен')
    macros = Macro.query.all()
    return render_template('macros.html', macros=macros)


@app.route('/macro_test', methods=['POST'])
@login_required
def macro_test():
    macro_type = request.form.get('macro_type')
    cfg = {}
    if macro_type == 'counter':
        cfg = {
            'start': int(request.form.get('start') or 0),
            'step': int(request.form.get('step') or 1),
            'current': int(request.form.get('start') or 0),
        }
    elif macro_type == 'random':
        cfg = {
            'chars': request.form.get('chars') or string.ascii_letters,
            'min_len': int(request.form.get('min_len') or 5),
            'max_len': int(request.form.get('max_len') or 10),
        }
    elif macro_type == 'list':
        uploaded = request.files.get('file')
        items = []
        if uploaded:
            text = uploaded.read().decode('utf-8')
            src = request.form.get('source') or 'lines'
            if src == 'words':
                items = re.findall(r'\w+', text)
            elif src == 'sentences':
                items = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
            else:
                items = [line.strip() for line in text.splitlines() if line.strip()]
        cfg = {
            'items': items,
            'mode': request.form.get('mode') or 'random',
            'index': 0,
            'source': request.form.get('source') or 'lines',
        }
        wmin = request.form.get('words_min')
        if wmin:
            cfg.update({
                'words_min': int(wmin or 1),
                'words_max': int(request.form.get('words_max') or wmin or 1),
                'sent_min': int(request.form.get('sent_min') or 1),
                'sent_max': int(request.form.get('sent_max') or 1),
                'para_min': int(request.form.get('para_min') or 1),
                'para_max': int(request.form.get('para_max') or 1),
                'as_html': bool(request.form.get('as_html')),
                'html_extra': request.form.get('html_extra') or '',
            })
    elif macro_type == 'multi':
        cfg = {
            'expr': request.form.get('expression') or '',
            'encoding': request.form.get('encoding') or 'none',
        }
    value = preview_macro_value(macro_type or '', cfg)
    return value


@app.route('/macros/delete/<int:macro_id>')
@login_required
def delete_macro(macro_id):
    macro = Macro.query.get_or_404(macro_id)
    db.session.delete(macro)
    db.session.commit()
    flash('Макрос удален')
    return redirect(url_for('macros'))


# ------- Attachments -------

def randomize_image(path: str, cfg: dict) -> None:
    img = Image.open(path).convert('RGB')
    left = random.randint(*cfg.get('left', (0, 0)))
    right = random.randint(*cfg.get('right', (0, 0)))
    top = random.randint(*cfg.get('top', (0, 0)))
    bottom = random.randint(*cfg.get('bottom', (0, 0)))
    new_w = img.width + left + right
    new_h = img.height + top + bottom
    new_img = Image.new('RGB', (new_w, new_h), 'white')
    new_img.paste(img, (left, top))
    draw = ImageDraw.Draw(new_img)
    colors = [
        '#000000', '#FFFFFF', '#FF0000', '#00FF00', '#0000FF', '#FFFF00', '#00FFFF',
        '#FF00FF', '#C0C0C0', '#808080', '#800000', '#808000', '#008000', '#800080',
        '#008080', '#000080'
    ]
    dot_min, dot_max = cfg.get('dots', (0, 0))
    for _ in range(random.randint(dot_min, dot_max)):
        x = random.randint(0, new_w - 1)
        y = random.randint(0, new_h - 1)
        draw.point((x, y), fill=random.choice(colors))
    new_img.save(path)


def random_attach_name() -> str:
    """Generate unique randomized filename base."""
    part1 = ''.join(random.choices(string.ascii_uppercase + string.digits, k=random.randint(8, 16)))
    part2 = ''.join(random.choices(string.ascii_uppercase + string.digits, k=random.randint(8, 16)))
    return f"{part1}_{part2}"
@app.route('/attachments', methods=['GET', 'POST'])
@login_required
def attachments():
    if request.method == 'POST':
        file = request.files.get('file')
        display_name = request.form.get('display_name') or (file.filename if file else '')
        send_name = request.form.get('send_filename') or '{{$const_attach_1_name_file}}'
        inline = bool(request.form.get('inline'))
        upload = bool(request.form.get('upload_to_server'))
        randomize = bool(request.form.get('randomize'))
        if file and file.filename:
            filename = secure_filename(file.filename)
            path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(path)
            ext = os.path.splitext(filename)[1]
            if '{{$const_attach_1_name_file}}' in send_name:
                send_name = random_attach_name() + ext
            else:
                send_name = secure_filename(send_name)
                if not os.path.splitext(send_name)[1]:
                    send_name += ext
            cfg = None
            if randomize and filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                cfg = {
                    'left': [int(request.form.get('pad_left_min') or 0), int(request.form.get('pad_left_max') or 0)],
                    'right': [int(request.form.get('pad_right_min') or 0), int(request.form.get('pad_right_max') or 0)],
                    'top': [int(request.form.get('pad_top_min') or 0), int(request.form.get('pad_top_max') or 0)],
                    'bottom': [int(request.form.get('pad_bottom_min') or 0), int(request.form.get('pad_bottom_max') or 0)],
                    'dots': [int(request.form.get('dot_min') or 0), int(request.form.get('dot_max') or 0)],
                }
                randomize_image(path, cfg)
            # DOCX specific processing
            if filename.lower().endswith('.docx'):
                pages_min = int(request.form.get('doc_pages_min') or 0)
                pages_max = int(request.form.get('doc_pages_max') or 0)
                page_text = request.form.get('doc_page_content') or ''
                if pages_max > 0:
                    try:
                        from docx import Document
                        doc = Document(path)
                        count = random.randint(pages_min, pages_max)
                        for _ in range(count):
                            doc.add_page_break()
                            doc.add_paragraph(render_macros(page_text))
                        doc.save(path)
                    except Exception:
                        pass
                if request.form.get('convert_to_pdf'):
                    try:
                        from docx import Document
                        from fpdf import FPDF
                        doc = Document(path)
                        text = '\n'.join(p.text for p in doc.paragraphs)
                        pdf_path = path.rsplit('.', 1)[0] + '.pdf'
                        pdf = FPDF()
                        pdf.set_auto_page_break(auto=True, margin=15)
                        pdf.add_page()
                        pdf.set_font('Arial', size=12)
                        for line in text.split('\n'):
                            pdf.multi_cell(0, 10, line)
                        author = render_macros(request.form.get('pdf_author') or '')
                        title = render_macros(request.form.get('pdf_title') or '')
                        if title:
                            pdf.set_title(title)
                        if author:
                            pdf.set_author(author)
                        pdf.output(pdf_path)
                        os.remove(path)
                        filename = os.path.basename(pdf_path)
                        path = pdf_path
                    except Exception:
                        pass
            attach = Attachment(display_name=display_name, filename=filename, send_filename=send_name,
                                path=path, inline=inline, upload_to_server=upload, config=cfg)
            db.session.add(attach)
            db.session.commit()
            attach.macro_url = f'url_attach_{attach.id}'
            attach.macro_id = f'id_attach_{attach.id}'
            if inline and filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif')):
                attach.macro_base64 = f'attach_img_{attach.id}_base64'
            db.session.commit()
            flash('Вложение добавлено')
    attachments = Attachment.query.all()
    return render_template('attachments.html', attachments=attachments)


@app.route('/attachments/delete/<int:att_id>')
@login_required
def delete_attachment(att_id):
    att = Attachment.query.get_or_404(att_id)
    try:
        if att.path and os.path.exists(att.path):
            os.remove(att.path)
    except Exception:
        pass
    db.session.delete(att)
    db.session.commit()
    flash('Вложение удалено')
    return redirect(url_for('attachments'))


# ------- API Rules -------
@app.route('/api_rules', methods=['GET', 'POST'])
@login_required
def api_rules():
    limit = get_setting('per_account_limit', '1')
    cycle = get_setting('cycle_accounts', 'no')
    attempts = get_setting('send_attempts', '1')
    timeout = get_setting('server_timeout', '30')
    pause = get_setting('pause_between', '0')
    recipients = get_setting('recipients_per_message', '1')
    method = get_setting('recipient_method', 'bcc')
    first = get_setting('first_recipient_to', 'no')
    q_every = get_setting('quality_every', '0')
    q_email = get_setting('quality_email', '')
    if request.method == 'POST':
        limit = request.form.get('per_account_limit') or '1'
        cycle = 'yes' if request.form.get('cycle_accounts') else 'no'
        attempts = request.form.get('send_attempts') or '1'
        timeout = request.form.get('server_timeout') or '30'
        pause = request.form.get('pause_between') or '0'
        recipients = request.form.get('recipients_per_message') or '1'
        method = request.form.get('recipient_method') or 'bcc'
        first = 'yes' if request.form.get('first_recipient_to') else 'no'
        q_every = request.form.get('quality_every') or '0'
        q_email = request.form.get('quality_email') or ''
        Setting.query.filter_by(key='per_account_limit').first().value = limit
        Setting.query.filter_by(key='cycle_accounts').first().value = cycle
        Setting.query.filter_by(key='send_attempts').first().value = attempts
        Setting.query.filter_by(key='server_timeout').first().value = timeout
        Setting.query.filter_by(key='pause_between').first().value = pause
        Setting.query.filter_by(key='recipients_per_message').first().value = recipients
        Setting.query.filter_by(key='recipient_method').first().value = method
        Setting.query.filter_by(key='first_recipient_to').first().value = first
        Setting.query.filter_by(key='quality_every').first().value = q_every
        Setting.query.filter_by(key='quality_email').first().value = q_email
        db.session.commit()
        flash('Правила сохранены')
    return render_template(
        'api_rules.html',
        limit=limit,
        cycle=cycle == 'yes',
        attempts=attempts,
        timeout=timeout,
        pause=pause,
        recipients=recipients,
        method=method,
        first=first == 'yes',
        q_every=q_every,
        q_email=q_email,
    )


# ------- Proxies -------
@app.route('/proxies', methods=['GET', 'POST'])
@login_required
def proxies():
    if request.method == 'POST':
        file = request.files.get('file')
        if file:
            lines = file.read().decode().splitlines()
            for line in lines:
                addr = line.strip()
                if addr and not Proxy.query.filter_by(address=addr).first():
                    db.session.add(Proxy(address=addr))
            db.session.commit()
            flash('Прокси загружены')
    proxies = Proxy.query.all()
    return render_template('proxies.html', proxies=proxies)


@app.route('/proxies/delete/<int:proxy_id>')
@login_required
def delete_proxy(proxy_id):
    proxy = Proxy.query.get_or_404(proxy_id)
    db.session.delete(proxy)
    db.session.commit()
    flash('Прокси удален')
    return redirect(url_for('proxies'))


# ------- API Accounts -------
@app.route('/api_accounts', methods=['GET', 'POST'])
@login_required
def api_accounts():
    if request.method == 'POST':
        file = request.files.get('file')
        if file:
            lines = file.read().decode().splitlines()
            for line in lines:
                parts = line.strip().split(':')
                if len(parts) >= 6:
                    login, password, first_name, last_name, api_key, uuid = parts[:6]
                    if not ApiAccount.query.filter_by(login=login).first():
                        db.session.add(ApiAccount(login=login, password=password,
                                                   first_name=first_name, last_name=last_name,
                                                   api_key=api_key, uuid=uuid))
            db.session.commit()
            flash('Аккаунты загружены')
    accounts = ApiAccount.query.all()
    return render_template('api_accounts.html', accounts=accounts)


@app.route('/api_accounts/delete/<int:acc_id>')
@login_required
def delete_api_account(acc_id):
    acc = ApiAccount.query.get_or_404(acc_id)
    db.session.delete(acc)
    db.session.commit()
    flash('Аккаунт удален')
    return redirect(url_for('api_accounts'))


@app.route('/api_accounts/reset_counts')
@login_required
def reset_counts():
    ApiAccount.query.update({ApiAccount.send_count: 0, ApiAccount.in_use: False})
    db.session.commit()
    flash('Счетчики сброшены')
    return redirect(url_for('api_accounts'))


# ------- Settings -------
@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    domain_setting = Setting.query.filter_by(key='domain').first()
    ua_setting = Setting.query.filter_by(key='user_agent').first()
    threads_setting = Setting.query.filter_by(key='threads').first()
    if request.method == 'POST':
        domain = request.form.get('domain')
        user_agent = request.form.get('user_agent')
        threads = request.form.get('threads')
        password = request.form.get('password')
        if domain:
            domain_setting.value = domain
        if user_agent and ua_setting:
            ua_setting.value = user_agent
        if threads and threads_setting:
            threads_setting.value = threads
        if password:
            user = User.query.filter_by(username='admin').first()
            user.set_password(password)
        db.session.commit()
        flash('Настройки сохранены')
    domain = domain_setting.value if domain_setting else ''
    user_agent = ua_setting.value if ua_setting else ''
    threads = threads_setting.value if threads_setting else ''
    return render_template('settings.html', domain=domain, user_agent=user_agent, threads=threads)


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
